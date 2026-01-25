"""DOCX parser using python-docx library."""

from pathlib import Path
from typing import Dict, List, Any, Optional

from src.ingestion.parsers.base_parser import BaseParser, ParserError
from src.config.logging_config import get_logger
from src.utils.file_utils import get_file_extension

logger = get_logger(__name__)

try:
    from docx import Document
    from docx.document import Document as DocumentType
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. DOCX parsing will not work.")


class DOCXParser(BaseParser):
    """DOCX parser using python-docx library."""

    def __init__(self, preserve_formatting: bool = False):
        """
        Initialize DOCX parser.

        Args:
            preserve_formatting: Whether to preserve formatting hints (bold, italic)
        """
        self.preserve_formatting = preserve_formatting

    def supports(self, file_path: Path) -> bool:
        """
        Check if the parser supports the given file type.

        Args:
            file_path: Path to the document file

        Returns:
            True if the parser supports this file type
        """
        return get_file_extension(file_path) in [".docx", ".doc"]

    def get_supported_extensions(self) -> List[str]:
        """
        Get list of file extensions supported by this parser.

        Returns:
            List of supported file extensions
        """
        return [".docx", ".doc"]

    def parse(self, file_path: Path) -> Dict[str, Any]:
        """
        Parse a DOCX file.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Dictionary containing parsed content and metadata:
            {
                "text": str,           # Main text content
                "tables": List[Dict],   # Extracted tables
                "metadata": Dict,       # Document metadata
                "pages": int,          # Number of pages (estimated)
                "parser_type": str     # Parser identifier
            }

        Raises:
            ParserError: If parsing fails
        """
        if not DOCX_AVAILABLE:
            raise ParserError(
                "python-docx library is not installed. Install it with: pip install python-docx"
            )

        try:
            logger.info(f"Parsing DOCX file: {file_path}")

            # Open document
            doc = Document(str(file_path))

            # Extract text content
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    if self.preserve_formatting:
                        formatted_text = self._extract_formatted_text(paragraph)
                        text_parts.append(formatted_text)
                    else:
                        text_parts.append(paragraph.text)

            text_content = "\n\n".join(text_parts)

            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = self._extract_table_data(table)
                if table_data:
                    tables.append(table_data)

            # Extract headers and footers
            headers_footers = self._extract_headers_footers(doc)

            # Extract metadata
            metadata = self.extract_metadata(file_path)
            core_props = doc.core_properties
            metadata.update({
                "title": core_props.title,
                "author": core_props.author,
                "subject": core_props.subject,
                "keywords": core_props.keywords,
                "created": core_props.created.isoformat() if core_props.created else None,
                "modified": core_props.modified.isoformat() if core_props.modified else None,
                "revision": core_props.revision,
                "has_headers": len(headers_footers.get("headers", [])) > 0,
                "has_footers": len(headers_footers.get("footers", [])) > 0,
            })

            # Estimate page count (rough estimate based on content length)
            estimated_pages = self._estimate_page_count(text_content)

            result = {
                "text": text_content,
                "tables": tables,
                "metadata": metadata,
                "pages": estimated_pages,
                "parser_type": "DOCXParser",
                "headers": headers_footers.get("headers", []),
                "footers": headers_footers.get("footers", []),
            }

            logger.info(
                f"Successfully parsed DOCX: {file_path} "
                f"({estimated_pages} estimated pages, {len(tables)} tables)"
            )
            return result

        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}", exc_info=True)
            raise ParserError(f"Failed to parse DOCX: {str(e)}") from e

    def _extract_formatted_text(self, paragraph: Paragraph) -> str:
        """
        Extract text with formatting hints.

        Args:
            paragraph: Paragraph object

        Returns:
            Text with formatting markers
        """
        text_parts = []
        for run in paragraph.runs:
            text = run.text
            if run.bold:
                text = f"**{text}**"
            if run.italic:
                text = f"*{text}*"
            text_parts.append(text)
        return "".join(text_parts) if text_parts else paragraph.text

    def _extract_table_data(self, table: Table) -> Optional[Dict[str, Any]]:
        """
        Extract data from a table.

        Args:
            table: Table object from python-docx

        Returns:
            Dictionary with table data or None
        """
        try:
            rows = []
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    cells.append(cell_text)
                rows.append(cells)

            if not rows:
                return None

            # Try to identify header row (first row with non-empty cells)
            header_row = None
            if rows and any(cell for cell in rows[0]):
                header_row = rows[0]
                data_rows = rows[1:]
            else:
                data_rows = rows

            return {
                "rows": rows,
                "header_row": header_row,
                "data_rows": data_rows,
                "row_count": len(rows),
                "column_count": len(rows[0]) if rows else 0,
                "type": "table",
            }
        except Exception as e:
            logger.warning(f"Could not extract table data: {e}")
            return None

    def _extract_headers_footers(self, doc: DocumentType) -> Dict[str, List[str]]:
        """
        Extract headers and footers from document.

        Args:
            doc: Document object

        Returns:
            Dictionary with headers and footers
        """
        headers = []
        footers = []

        try:
            for section in doc.sections:
                # Extract header
                if section.header:
                    header_texts = []
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            header_texts.append(paragraph.text)
                    if header_texts:
                        headers.extend(header_texts)

                # Extract footer
                if section.footer:
                    footer_texts = []
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            footer_texts.append(paragraph.text)
                    if footer_texts:
                        footers.extend(footer_texts)
        except Exception as e:
            logger.warning(f"Could not extract headers/footers: {e}")

        return {
            "headers": list(set(headers)),  # Remove duplicates
            "footers": list(set(footers)),
        }

    def _estimate_page_count(self, text: str) -> int:
        """
        Estimate page count based on text length.

        Args:
            text: Document text content

        Returns:
            Estimated number of pages
        """
        # Rough estimate: ~500 words per page, ~5 characters per word
        # This is a very rough estimate
        char_count = len(text)
        words_per_page = 500
        chars_per_word = 5
        chars_per_page = words_per_page * chars_per_word

        estimated_pages = max(1, char_count // chars_per_page)
        return estimated_pages

